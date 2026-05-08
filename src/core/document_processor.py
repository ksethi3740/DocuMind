"""
Handles PDF, DOCX, and image ingestion.
Extracts text, tables, and embedded images.
"""

import io
import fitz                    # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document
from pathlib import Path
from langchain_text_splitters import RecursiveCharacterTextSplitter
from config.settings import CHUNK_SIZE, CHUNK_OVERLAP


def extract_text_from_pdf(file_bytes: bytes) -> dict:
    """Extract text and images from PDF using PyMuPDF."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages_text = []
    images_text = []

    for page_num, page in enumerate(doc):
        # Full text
        text = page.get_text("text")
        pages_text.append({"page": page_num + 1, "text": text})

        # Embedded images → OCR
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            pil_img = Image.open(io.BytesIO(image_bytes))
            try:
                ocr_text = pytesseract.image_to_string(pil_img)
                if ocr_text.strip():
                    images_text.append({
                        "page": page_num + 1,
                        "image_index": img_index,
                        "text": ocr_text
                    })
            except Exception:
                pass

    doc.close()
    full_text = "\n\n".join(p["text"] for p in pages_text if p["text"].strip())
    image_ocr = "\n\n".join(i["text"] for i in images_text if i["text"].strip())

    return {
        "full_text": full_text,
        "image_ocr": image_ocr,
        "pages": pages_text,
        "page_count": len(pages_text)
    }


def extract_text_from_docx(file_bytes: bytes) -> dict:
    """Extract text from DOCX including tables."""
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    table_data = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                table_data.append(row_text)

    full_text = "\n".join(paragraphs)
    if table_data:
        full_text += "\n\n[TABLES]\n" + "\n".join(table_data)

    return {"full_text": full_text, "paragraph_count": len(paragraphs)}


def extract_text_from_image(file_bytes: bytes) -> dict:
    """OCR an uploaded image file."""
    pil_img = Image.open(io.BytesIO(file_bytes))
    text = pytesseract.image_to_string(pil_img)
    return {"full_text": text.strip(), "source": "ocr"}


def process_uploaded_file(file_obj) -> dict:
    """
    Universal file processor.
    Returns: {name, text, metadata, chunks}
    """
    file_bytes = file_obj.read()
    name = file_obj.name.lower()

    if name.endswith(".pdf"):
        data = extract_text_from_pdf(file_bytes)
        text = data["full_text"]
        if data["image_ocr"]:
            text += "\n\n[Image OCR]\n" + data["image_ocr"]
        meta = {"type": "pdf", "pages": data["page_count"]}

    elif name.endswith(".docx"):
        data = extract_text_from_docx(file_bytes)
        text = data["full_text"]
        meta = {"type": "docx", "paragraphs": data["paragraph_count"]}

    elif name.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp")):
        data = extract_text_from_image(file_bytes)
        text = data["full_text"]
        meta = {"type": "image"}

    else:
        text = file_bytes.decode("utf-8", errors="ignore")
        meta = {"type": "text"}

    # Split into chunks
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " "]
    )
    chunks = splitter.split_text(text)

    return {
        "name": file_obj.name,
        "text": text,
        "chunks": chunks,
        "metadata": meta,
        "chunk_count": len(chunks)
    }